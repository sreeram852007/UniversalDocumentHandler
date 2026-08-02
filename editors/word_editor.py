"""
Word Editor - Microsoft Word-like rich text editor with formatting toolbar and zoom support
"""

import os
import re
from PyQt6.QtWidgets import *
from PyQt6.QtCore import *
from PyQt6.QtGui import *
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ui.styles import ThemeManager


class WordEditor(QWidget):
    """Full-featured Word editor with Microsoft Word-like ribbon UI"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.current_file = None
        self.is_modified = False
        self.zoom_factor = 1.0
        self.is_fullscreen = False
        self.text_edit = None
        self.bold_btn = None
        self.italic_btn = None
        self.underline_btn = None
        self.font_combo = None
        self.size_combo = None
        self.zoom_combo = None
        self.color_btn = None
        self.word_count_label = None
        self.page_info_label = None
        self.scroll_area = None
        self.init_ui()
        self.setup_shortcuts()
    
    def init_ui(self):
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        self.setStyleSheet(ThemeManager.get_editor_style("word"))
        
        # === MS Word-like Ribbon ===
        ribbon = QToolBar()
        ribbon.setMovable(False)
        ribbon.setStyleSheet("""
            QToolBar {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #ffffff, stop:1 #f0f1f3);
                border: none;
                border-bottom: 2px solid #0078d4;
                padding: 2px 5px;
                min-height: 42px;
                spacing: 2px;
            }
            QToolBar::separator {
                width: 2px;
                background: #dee2e6;
                margin: 5px 3px;
            }
            QToolBar QToolButton {
                background-color: transparent;
                padding: 4px 8px;
                border-radius: 4px;
                font-weight: 500;
                font-size: 12px;
                color: #2c3e50;
            }
            QToolBar QToolButton:hover {
                background-color: #e8eaed;
            }
            QToolBar QToolButton:checked {
                background-color: #cce5ff;
                border: 1px solid #0078d4;
            }
            QComboBox {
                padding: 3px 8px;
                border: 1px solid #ced4da;
                border-radius: 4px;
                background-color: white;
                font-size: 12px;
                min-height: 24px;
            }
            QComboBox:hover {
                border-color: #0078d4;
            }
            QPushButton {
                background-color: #0078d4;
                color: white;
                border: none;
                padding: 4px 10px;
                border-radius: 3px;
                font-weight: 600;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #106ebe;
            }
        """)
        
        # === Font Group ===
        ribbon.addWidget(QLabel("Font:"))
        self.font_combo = QComboBox()
        self.font_combo.addItems([
            "Arial", "Calibri", "Cambria", "Candara", "Consolas",
            "Courier New", "Georgia", "Impact", "Segoe UI", "Times New Roman",
            "Trebuchet MS", "Verdana"
        ])
        self.font_combo.setCurrentText("Segoe UI")
        self.font_combo.setMinimumWidth(120)
        self.font_combo.currentTextChanged.connect(self.change_font)
        ribbon.addWidget(self.font_combo)
        
        ribbon.addWidget(QLabel("Size:"))
        self.size_combo = QComboBox()
        self.size_combo.addItems([str(i) for i in [8, 9, 10, 11, 12, 14, 16, 18, 20, 24, 28, 32, 36, 48, 72]])
        self.size_combo.setCurrentText("12")
        self.size_combo.setMaximumWidth(60)
        self.size_combo.currentTextChanged.connect(self.change_font_size)
        ribbon.addWidget(self.size_combo)
        
        ribbon.addSeparator()
        
        # === Formatting Buttons ===
        self.bold_btn = QAction("B", self)
        self.bold_btn.setCheckable(True)
        self.bold_btn.triggered.connect(self.toggle_bold)
        self.bold_btn.setShortcut("Ctrl+B")
        self.bold_btn.setToolTip("Bold (Ctrl+B)")
        self.bold_btn.setFont(QFont("Segoe UI", 10, QFont.Weight.Bold))
        ribbon.addAction(self.bold_btn)
        
        self.italic_btn = QAction("I", self)
        self.italic_btn.setCheckable(True)
        self.italic_btn.triggered.connect(self.toggle_italic)
        self.italic_btn.setShortcut("Ctrl+I")
        self.italic_btn.setToolTip("Italic (Ctrl+I)")
        self.italic_btn.setFont(QFont("Segoe UI", 10, QFont.Weight.Normal, True))
        ribbon.addAction(self.italic_btn)
        
        self.underline_btn = QAction("U", self)
        self.underline_btn.setCheckable(True)
        self.underline_btn.triggered.connect(self.toggle_underline)
        self.underline_btn.setShortcut("Ctrl+U")
        self.underline_btn.setToolTip("Underline (Ctrl+U)")
        ribbon.addAction(self.underline_btn)
        
        ribbon.addSeparator()
        
        # === Color ===
        self.color_btn = QPushButton("Color")
        self.color_btn.clicked.connect(self.choose_color)
        self.color_btn.setMaximumWidth(50)
        self.color_btn.setStyleSheet("background-color: #000000; color: white; border: 1px solid #ccc; border-radius: 3px;")
        ribbon.addWidget(self.color_btn)
        
        ribbon.addSeparator()
        
        # === Alignment ===
        align_left = QAction("⇐", self)
        align_left.triggered.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignLeft))
        ribbon.addAction(align_left)
        
        align_center = QAction("⇔", self)
        align_center.triggered.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignCenter))
        ribbon.addAction(align_center)
        
        align_right = QAction("⇒", self)
        align_right.triggered.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignRight))
        ribbon.addAction(align_right)
        
        align_justify = QAction("≡", self)
        align_justify.triggered.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignJustify))
        ribbon.addAction(align_justify)
        
        ribbon.addSeparator()
        
        # === Lists ===
        bullet_btn = QAction("•", self)
        bullet_btn.triggered.connect(self.toggle_bullet_list)
        ribbon.addAction(bullet_btn)
        
        number_btn = QAction("1.", self)
        number_btn.triggered.connect(self.toggle_numbered_list)
        ribbon.addAction(number_btn)
        
        ribbon.addSeparator()
        
        # === Undo/Redo ===
        undo_btn = QAction("↩ Undo", self)
        undo_btn.setShortcut("Ctrl+Z")
        undo_btn.triggered.connect(self.undo)
        ribbon.addAction(undo_btn)
        
        redo_btn = QAction("↪ Redo", self)
        redo_btn.setShortcut("Ctrl+Y")
        redo_btn.triggered.connect(self.redo)
        ribbon.addAction(redo_btn)
        
        ribbon.addSeparator()
        
        # === Insert ===
        insert_image_btn = QAction("🖼️ Image", self)
        insert_image_btn.triggered.connect(self.insert_image)
        ribbon.addAction(insert_image_btn)
        
        ribbon.addSeparator()
        
        # === Save ===
        save_action = QAction("💾 Save", self)
        save_action.setShortcut("Ctrl+S")
        save_action.triggered.connect(self.save_document)
        ribbon.addAction(save_action)
        
        layout.addWidget(ribbon)
        
        # === Zoom Toolbar ===
        zoom_toolbar = QToolBar()
        zoom_toolbar.setMovable(False)
        zoom_toolbar.setStyleSheet("""
            QToolBar {
                background-color: #f8f9fa;
                border: 1px solid #dee2e6;
                padding: 2px;
                min-height: 32px;
            }
            QToolBar QToolButton {
                background-color: transparent;
                padding: 2px 8px;
                border-radius: 3px;
            }
            QToolBar QToolButton:hover {
                background-color: #e8eaed;
            }
            QComboBox {
                padding: 2px 8px;
                border: 1px solid #ced4da;
                border-radius: 3px;
                background-color: white;
                font-size: 12px;
            }
        """)
        
        zoom_out_btn = QAction("🔍-", self)
        zoom_out_btn.triggered.connect(self.zoom_out)
        zoom_out_btn.setShortcut("Ctrl+-")
        zoom_toolbar.addAction(zoom_out_btn)
        
        self.zoom_combo = QComboBox()
        self.zoom_combo.addItems(["50%", "75%", "100%", "125%", "150%", "200%", "300%"])
        self.zoom_combo.setCurrentText("100%")
        self.zoom_combo.setMaximumWidth(80)
        self.zoom_combo.currentTextChanged.connect(self.zoom_changed)
        zoom_toolbar.addWidget(self.zoom_combo)
        
        zoom_in_btn = QAction("🔍+", self)
        zoom_in_btn.triggered.connect(self.zoom_in)
        zoom_in_btn.setShortcut("Ctrl++")
        zoom_toolbar.addAction(zoom_in_btn)
        
        zoom_toolbar.addSeparator()
        
        fit_width_btn = QAction("Fit to Width", self)
        fit_width_btn.triggered.connect(self.fit_to_width)
        zoom_toolbar.addAction(fit_width_btn)
        
        fullscreen_btn = QAction("⛶ Full Screen", self)
        fullscreen_btn.triggered.connect(self.toggle_fullscreen)
        fullscreen_btn.setShortcut("F11")
        zoom_toolbar.addAction(fullscreen_btn)
        
        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        zoom_toolbar.addWidget(spacer)
        
        self.word_count_label = QLabel("Words: 0")
        zoom_toolbar.addWidget(self.word_count_label)
        
        layout.addWidget(zoom_toolbar)
        
        # === Main Editor ===
        self.text_edit = QTextEdit()
        self.text_edit.setFont(QFont("Segoe UI", 12))
        self.text_edit.setStyleSheet("""
            QTextEdit {
                background-color: white;
                border: none;
                padding: 40px 50px;
                font-family: 'Segoe UI', Arial, sans-serif;
                font-size: 12pt;
                line-height: 1.6;
            }
            QTextEdit:focus {
                border: none;
            }
        """)
        self.text_edit.textChanged.connect(self.on_text_changed)
        self.text_edit.setFocus()
        self.text_edit.wheelEvent = self.zoom_wheel_event
        
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setStyleSheet("background-color: #e8e8e8; border: none;")
        self.scroll_area.setWidget(self.text_edit)
        layout.addWidget(self.scroll_area)
        
        # === Status Bar ===
        status_bar = QStatusBar()
        status_bar.setStyleSheet("""
            QStatusBar {
                background-color: #f8f9fa;
                border-top: 1px solid #dee2e6;
                padding: 2px 10px;
                font-size: 12px;
                color: #495057;
            }
        """)
        self.page_info_label = QLabel("Lines: 0 | Characters: 0 | Words: 0 | Zoom: 100%")
        status_bar.addWidget(self.page_info_label)
        layout.addWidget(status_bar)
        
        self.setLayout(layout)
        self.update_status()
    
    def setup_shortcuts(self):
        """Setup keyboard shortcuts"""
        if self.text_edit:
            self.text_edit.wheelEvent = self.zoom_wheel_event
    
    def zoom_wheel_event(self, event):
        """Handle zoom with Ctrl+Scroll"""
        if event.modifiers() & Qt.KeyboardModifier.ControlModifier:
            delta = event.angleDelta().y()
            if delta > 0:
                self.zoom_in()
            else:
                self.zoom_out()
            event.accept()
        else:
            QTextEdit.wheelEvent(self.text_edit, event)
    
    def zoom_changed(self, value):
        """Handle zoom combo box changes"""
        self.zoom_factor = float(value.replace("%", "")) / 100
        self.apply_zoom()
        self.update_status()
    
    def zoom_in(self):
        """Zoom in"""
        current = self.zoom_combo.currentText()
        zoom_values = ["50%", "75%", "100%", "125%", "150%", "200%", "300%"]
        if current in zoom_values:
            idx = zoom_values.index(current)
            if idx < len(zoom_values) - 1:
                self.zoom_combo.setCurrentText(zoom_values[idx + 1])
                self.update_status()
    
    def zoom_out(self):
        """Zoom out"""
        current = self.zoom_combo.currentText()
        zoom_values = ["50%", "75%", "100%", "125%", "150%", "200%", "300%"]
        if current in zoom_values:
            idx = zoom_values.index(current)
            if idx > 0:
                self.zoom_combo.setCurrentText(zoom_values[idx - 1])
                self.update_status()
    
    def apply_zoom(self):
        """Apply zoom factor to text"""
        base_size = 12
        new_size = base_size * self.zoom_factor
        self.text_edit.setFontPointSize(new_size)
    
    def fit_to_width(self):
        """Fit text to window width"""
        self.zoom_combo.setCurrentText("100%")
        self.update_status()
    
    def toggle_fullscreen(self):
        """Toggle full-screen mode"""
        parent = self.parent()
        while parent and not isinstance(parent, QMainWindow):
            parent = parent.parent()
        
        if parent:
            self.is_fullscreen = not self.is_fullscreen
            if self.is_fullscreen:
                parent.showFullScreen()
            else:
                parent.showNormal()
    
    def change_font(self, font_name):
        self.text_edit.setFontFamily(font_name)
    
    def change_font_size(self, size):
        base_size = int(size)
        actual_size = base_size * self.zoom_factor
        self.text_edit.setFontPointSize(actual_size)
    
    def toggle_bold(self):
        if self.bold_btn.isChecked():
            self.text_edit.setFontWeight(QFont.Weight.Bold)
        else:
            self.text_edit.setFontWeight(QFont.Weight.Normal)
    
    def toggle_italic(self):
        self.text_edit.setFontItalic(self.italic_btn.isChecked())
    
    def toggle_underline(self):
        self.text_edit.setFontUnderline(self.underline_btn.isChecked())
    
    def choose_color(self):
        color = QColorDialog.getColor(Qt.GlobalColor.black, self, "Choose Text Color")
        if color.isValid():
            self.text_edit.setTextColor(color)
            self.color_btn.setStyleSheet(f"background-color: {color.name()}; color: {'white' if color.lightness() < 128 else 'black'}; border: 1px solid #ccc; border-radius: 3px;")
    
    def set_alignment(self, alignment):
        self.text_edit.setAlignment(alignment)
    
    def toggle_bullet_list(self):
        cursor = self.text_edit.textCursor()
        cursor.insertText("• ")
    
    def toggle_numbered_list(self):
        cursor = self.text_edit.textCursor()
        # Simple numbered list - get next number
        text = self.text_edit.toPlainText()
        lines = text.split('\n')
        number = 1
        for line in reversed(lines):
            if re.match(r'^\d+\.', line.strip()):
                try:
                    number = int(re.match(r'^(\d+)\.', line.strip()).group(1)) + 1
                    break
                except:
                    pass
        cursor.insertText(f"{number}. ")
    
    def insert_image(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Image",
            "",
            "Images (*.png *.jpg *.jpeg *.gif *.bmp)"
        )
        if file_path:
            cursor = self.text_edit.textCursor()
            cursor.insertImage(file_path)
    
    def on_text_changed(self):
        self.is_modified = True
        self.update_status()
    
    def update_status(self):
        """Update status bar with document info and zoom level"""
        if not self.text_edit:
            return
        
        text = self.text_edit.toPlainText()
        words = len(text.split())
        chars = len(text)
        lines = text.count('\n') + 1
        
        # Get current zoom percentage
        zoom_percent = int(self.zoom_factor * 100)
        
        if self.page_info_label:
            self.page_info_label.setText(f"Lines: {lines} | Characters: {chars} | Words: {words} | Zoom: {zoom_percent}%")
        
        # Update word count in toolbar
        if self.word_count_label:
            self.word_count_label.setText(f"Words: {words}")
    
    def undo(self):
        """Undo last action"""
        if self.text_edit:
            self.text_edit.undo()
    
    def redo(self):
        """Redo last action"""
        if self.text_edit:
            self.text_edit.redo()
    
    def load_document(self, file_path):
        """Load DOCX file into editor with formatting"""
        self.current_file = file_path
        try:
            doc = Document(file_path)
            
            html_parts = []
            
            # Process paragraphs with formatting
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                
                # Build formatted text
                para_html = self._paragraph_to_html(para)
                if para_html:
                    html_parts.append(para_html)
            
            # Process tables
            for table in doc.tables:
                table_html = self._table_to_html(table)
                if table_html:
                    html_parts.append(table_html)
            
            # Wrap in HTML
            html_content = self._wrap_html(html_parts)
            self.text_edit.setHtml(html_content)
            
            self.is_modified = False
            self.text_edit.document().setModified(False)
            self.update_status()
            
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to load document: {str(e)}")
    
    def _paragraph_to_html(self, para):
        """Convert a paragraph to HTML with formatting"""
        if not para.text.strip():
            return ""
        
        # Get paragraph style
        style_name = para.style.name if para.style else ""
        
        # Determine heading level
        heading_level = 0
        if "Heading" in style_name:
            try:
                heading_level = int(style_name.replace("Heading", "").strip())
            except:
                heading_level = 1
        
        # Build paragraph with formatting
        text_parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            # Apply formatting
            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"
            
            text_parts.append(text)
        
        full_text = "".join(text_parts)
        
        # Apply heading or paragraph styling
        if heading_level and heading_level <= 6:
            return f"<h{heading_level}>{full_text}</h{heading_level}>"
        else:
            # Check if it's a list item
            if para.text.startswith("•") or para.text.startswith("-") or para.text.startswith("◦"):
                return f'<li style="margin-left: 20px;">{full_text}</li>'
            
            # Check for numbering
            if re.match(r'^\d+\.', para.text.strip()):
                return f'<li style="margin-left: 20px; list-style-type: decimal;">{full_text}</li>'
            
            # Apply alignment
            alignment = para.alignment
            align_style = ""
            if alignment:
                if alignment == 1:  # Center
                    align_style = ' style="text-align: center;"'
                elif alignment == 2:  # Right
                    align_style = ' style="text-align: right;"'
                elif alignment == 3:  # Justify
                    align_style = ' style="text-align: justify;"'
            
            return f'<p{align_style}>{full_text}</p>'
    
    def _table_to_html(self, table):
        """Convert a table to HTML"""
        html = ['<table style="border-collapse: collapse; width: 100%; margin: 10px 0;">']
        
        for row in table.rows:
            html.append('<tr>')
            for cell in row.cells:
                cell_text = cell.text.strip()
                is_header = row == table.rows[0]
                tag = 'th' if is_header else 'td'
                style = 'border: 1px solid #ddd; padding: 8px; text-align: left;'
                if is_header:
                    style += ' background-color: #f2f2f2; font-weight: bold;'
                html.append(f'<{tag} style="{style}">{cell_text}</{tag}>')
            html.append('</tr>')
        
        html.append('</table>')
        return "\n".join(html)
    
    def _wrap_html(self, content_parts):
        """Wrap content in HTML document"""
        return f"""
        <html>
        <head>
            <style>
                body {{ 
                    font-family: 'Segoe UI', Arial, sans-serif;
                    font-size: 14px;
                    line-height: 1.6;
                    padding: 20px;
                    max-width: 900px;
                    margin: 0 auto;
                    background-color: #ffffff;
                }}
                h1 {{ font-size: 24px; color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                h2 {{ font-size: 20px; color: #2c3e50; margin-top: 20px; }}
                h3 {{ font-size: 18px; color: #2c3e50; }}
                h4 {{ font-size: 16px; color: #2c3e50; }}
                h5 {{ font-size: 14px; color: #2c3e50; }}
                h6 {{ font-size: 13px; color: #2c3e50; }}
                p {{ margin: 8px 0; }}
                ul, ol {{ padding-left: 30px; }}
                li {{ margin: 4px 0; }}
                table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
                th {{ background-color: #f2f2f2; font-weight: bold; }}
                td, th {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                tr:nth-child(even) {{ background-color: #f9f9f9; }}
                .highlight {{ background-color: #fff3cd; padding: 2px 4px; }}
                .note {{ background-color: #d1ecf1; padding: 10px; border-radius: 4px; margin: 10px 0; }}
            </style>
        </head>
        <body>
            {''.join(content_parts)}
        </body>
        </html>
        """
    
    def save_document(self):
        """Save the document as DOCX"""
        if not self.current_file:
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save Document",
                "Untitled.docx",
                "Word Documents (*.docx);;PDF Files (*.pdf);;Text Files (*.txt)"
            )
            if not file_path:
                return
            self.current_file = file_path
        
        try:
            doc = Document()
            text = self.text_edit.toPlainText()
            for line in text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            doc.save(self.current_file)
            self.is_modified = False
            self.text_edit.document().setModified(False)
            QMessageBox.information(self, "Success", "Document saved successfully!")
            
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to save: {str(e)}")
    
    def keyPressEvent(self, event):
        """Handle keyboard shortcuts"""
        if event.key() == Qt.Key.Key_Escape and self.is_fullscreen:
            self.toggle_fullscreen()
            event.accept()
        else:
            super().keyPressEvent(event)
    
    def resizeEvent(self, event):
        """Handle resize events"""
        super().resizeEvent(event)
        self.update_status()