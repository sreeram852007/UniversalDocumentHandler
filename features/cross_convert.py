"""
Cross-Format Conversion - Convert between ANY supported formats
"""

import os
import tempfile
from pathlib import Path
from typing import List, Dict, Optional
import pandas as pd
from docx import Document
import fitz  # PyMuPDF
from PIL import Image


class CrossFormatConverter:
    """Universal cross-format document converter"""
    
    SUPPORTED_FORMATS = {
        'pdf': ['PDF Document', '.pdf'],
        'docx': ['Word Document', '.docx'],
        'txt': ['Text File', '.txt'],
        'csv': ['CSV File', '.csv'],
        'xlsx': ['Excel Spreadsheet', '.xlsx'],
        'pptx': ['PowerPoint Presentation', '.pptx'],
        'jpg': ['JPEG Image', '.jpg'],
        'png': ['PNG Image', '.png'],
        'html': ['HTML Document', '.html'],
        'md': ['Markdown', '.md'],
        'epub': ['EPUB E-Book', '.epub'],
        'rtf': ['RTF Document', '.rtf'],
        'json': ['JSON Data', '.json'],
        'xml': ['XML Data', '.xml'],
    }
    
    def __init__(self):
        self.conversion_matrix = self._build_conversion_matrix()
    
    def _build_conversion_matrix(self) -> Dict:
        """Build matrix of supported conversions"""
        matrix = {}
        
        # Define supported conversions (from -> to)
        conversions = {
            'pdf': ['docx', 'txt', 'html', 'epub', 'jpg', 'png'],
            'docx': ['pdf', 'txt', 'html', 'rtf', 'md'],
            'txt': ['pdf', 'docx', 'html', 'md', 'json'],
            'csv': ['xlsx', 'pdf', 'html', 'json'],
            'xlsx': ['csv', 'pdf', 'html', 'json'],
            'pptx': ['pdf', 'jpg', 'png'],
            'jpg': ['png', 'pdf', 'html'],
            'png': ['jpg', 'pdf', 'html'],
            'html': ['pdf', 'docx', 'txt', 'md'],
            'md': ['pdf', 'html', 'docx'],
            'epub': ['pdf', 'txt', 'html'],
            'rtf': ['pdf', 'docx', 'txt'],
            'json': ['csv', 'xlsx', 'xml'],
            'xml': ['json', 'csv', 'xlsx'],
        }
        
        for fmt, targets in conversions.items():
            for target in targets:
                if fmt in self.SUPPORTED_FORMATS and target in self.SUPPORTED_FORMATS:
                    matrix[(fmt, target)] = True
        
        return matrix
    
    def get_supported_conversions(self, input_format: str) -> List[str]:
        """Get list of formats this input can convert to"""
        results = []
        for (in_fmt, out_fmt) in self.conversion_matrix:
            if in_fmt == input_format and out_fmt not in results:
                results.append(out_fmt)
        return results
    
    def convert(self, input_path: str, output_path: str, input_format: str, output_format: str) -> bool:
        """
        Convert file from one format to another
        
        Args:
            input_path: Source file path
            output_path: Destination file path
            input_format: Input format (pdf, docx, txt, etc.)
            output_format: Output format (pdf, docx, txt, etc.)
        
        Returns:
            True if successful, False otherwise
        """
        try:
            # Check if conversion is supported
            if (input_format, output_format) not in self.conversion_matrix:
                raise ValueError(f"Conversion from {input_format} to {output_format} not supported")
            
            # Dispatch to specific converter
            converter_method = f"_convert_{input_format}_to_{output_format}"
            
            if hasattr(self, converter_method):
                return getattr(self, converter_method)(input_path, output_path)
            else:
                # Use generic conversion
                return self._convert_generic(input_path, output_path, input_format, output_format)
                
        except Exception as e:
            print(f"Conversion error: {e}")
            return False
    
    def _convert_pdf_to_docx(self, input_path: str, output_path: str) -> bool:
        """Convert PDF to DOCX"""
        try:
            from pdf2docx import Converter
            cv = Converter(input_path)
            cv.convert(output_path)
            cv.close()
            return True
        except:
            # Fallback: extract text to DOCX
            doc = fitz.open(input_path)
            word_doc = Document()
            for page in doc:
                text = page.get_text()
                if text.strip():
                    word_doc.add_paragraph(text)
            word_doc.save(output_path)
            doc.close()
            return True
    
    def _convert_pdf_to_txt(self, input_path: str, output_path: str) -> bool:
        """Convert PDF to TXT"""
        doc = fitz.open(input_path)
        text = ""
        for page in doc:
            text += page.get_text() + "\n\n"
        doc.close()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        return True
    
    def _convert_pdf_to_jpg(self, input_path: str, output_path: str) -> bool:
        """Convert PDF to JPG"""
        doc = fitz.open(input_path)
        page = doc[0]
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        pix.save(output_path)
        doc.close()
        return True
    
    def _convert_docx_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert DOCX to PDF"""
        doc = Document(input_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        page.insert_text((50, 50), text[:5000])
        pdf_doc.save(output_path)
        pdf_doc.close()
        return True
    
    def _convert_csv_to_xlsx(self, input_path: str, output_path: str) -> bool:
        """Convert CSV to XLSX"""
        df = pd.read_csv(input_path)
        df.to_excel(output_path, index=False)
        return True
    
    def _convert_xlsx_to_csv(self, input_path: str, output_path: str) -> bool:
        """Convert XLSX to CSV"""
        df = pd.read_excel(input_path)
        df.to_csv(output_path, index=False)
        return True
    
    def _convert_image_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert Image to PDF"""
        img = Image.open(input_path)
        if img.mode in ('RGBA', 'LA', 'P'):
            img = img.convert('RGB')
        img.save(output_path, 'PDF', resolution=100.0)
        return True
    
    def _convert_generic(self, input_path: str, output_path: str, input_format: str, output_format: str) -> bool:
        """Generic conversion via text extraction"""
        try:
            # Extract text based on input format
            text = self._extract_text(input_path, input_format)
            
            if not text:
                return False
            
            # Write text based on output format
            if output_format == 'txt':
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
            elif output_format == 'docx':
                doc = Document()
                for line in text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
                doc.save(output_path)
            elif output_format == 'html':
                html = f"<html><body><pre>{text}</pre></body></html>"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html)
            elif output_format == 'pdf':
                pdf_doc = fitz.open()
                page = pdf_doc.new_page()
                page.insert_text((50, 50), text[:5000])
                pdf_doc.save(output_path)
                pdf_doc.close()
            else:
                return False
            
            return True
            
        except Exception as e:
            print(f"Generic conversion failed: {e}")
            return False
    
    def _extract_text(self, file_path: str, file_format: str) -> str:
        """Extract text from various file formats"""
        try:
            if file_format == 'pdf':
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
                return text
            
            elif file_format == 'docx':
                doc = Document(file_path)
                return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
            elif file_format in ['csv', 'xlsx']:
                if file_format == 'csv':
                    df = pd.read_csv(file_path)
                else:
                    df = pd.read_excel(file_path)
                return df.to_string()
            
            elif file_format == 'txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            
            elif file_format == 'html':
                import re
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    html = f.read()
                return re.sub(r'<[^>]+>', '', html)
            
            elif file_format == 'json':
                import json
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                return json.dumps(data, indent=2)
            
            else:
                # Try reading as text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
        except Exception as e:
            print(f"Text extraction failed: {e}")
            return ""
    
    def get_format_info(self, format_name: str) -> Optional[Dict]:
        """Get information about a format"""
        if format_name in self.SUPPORTED_FORMATS:
            name, ext = self.SUPPORTED_FORMATS[format_name]
            return {
                'name': name,
                'extension': ext,
                'format': format_name
            }
        return None
    
    def get_all_formats(self) -> List[str]:
        """Get list of all supported formats"""
        return list(self.SUPPORTED_FORMATS.keys())
    
    def get_display_name(self, format_name: str) -> str:
        """Get display name for a format"""
        if format_name in self.SUPPORTED_FORMATS:
            return self.SUPPORTED_FORMATS[format_name][0]
        return format_name.upper()